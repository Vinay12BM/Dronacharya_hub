from docx import Document
from docx.shared import Pt
import re

def create_docx(markdown_content, filepath):
    doc = Document()
    
    # Simple markdown to DOCX conversion
    # Process only headers and paragraphs/bullets for now
    lines = markdown_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        # Headers
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        # Bullets
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        # Simple Bold/Italic cleaning
        else:
            cleaned = re.sub(r'\*\*|\*', '', line)
            doc.add_paragraph(cleaned)
            
    try:
        doc.save(filepath)
        return True
    except Exception as e:
        print(f"DOCX Error: {e}")
        return False
